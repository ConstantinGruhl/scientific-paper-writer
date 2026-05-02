from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image,
    ListFlowable,
    ListItem,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table as PdfTable,
    TableStyle,
)

from .evidence_assets import figure_caption_text, lookup_figure, lookup_table, table_display_model
from .manuscript import render_document_model
from .profiles import resolve_effective_formatting_profile
from .references import bibliography_entries, build_citation_order, source_index
from .validation import validate_project_state
from .workspace import append_trace, load_all_states, save_state


def _page_size(profile: dict[str, Any]) -> tuple[float, float]:
    size_name = profile.get("page_size", "A4").upper()
    return LETTER if size_name == "LETTER" else A4


def _docx_page_size(profile: dict[str, Any]) -> tuple[float, float]:
    size_name = profile.get("page_size", "A4").upper()
    if size_name == "LETTER":
        return Inches(8.5), Inches(11.0)
    return Inches(8.27), Inches(11.69)


def _reportlab_font_name(font_name: str) -> str:
    mapping = {
        "times new roman": "Times-Roman",
        "times": "Times-Roman",
        "arial": "Helvetica",
        "helvetica": "Helvetica",
        "courier new": "Courier",
        "courier": "Courier",
    }
    return mapping.get(font_name.strip().lower(), "Times-Roman")


def _apply_docx_page_layout(document: Document, formatting_profile: dict[str, Any]) -> None:
    section = document.sections[0]
    page_width, page_height = _docx_page_size(formatting_profile)
    section.page_width = page_width
    section.page_height = page_height
    margins = formatting_profile.get("margins_inches", {})
    section.top_margin = Inches(margins.get("top", 1.0))
    section.bottom_margin = Inches(margins.get("bottom", 1.0))
    section.left_margin = Inches(margins.get("left", 1.0))
    section.right_margin = Inches(margins.get("right", 1.0))


def _configure_docx_styles(document: Document, formatting_profile: dict[str, Any]) -> None:
    font_name = formatting_profile.get("font_name", "Times New Roman")
    font_size = formatting_profile.get("font_size_pt", 12)
    line_spacing = formatting_profile.get("line_spacing", 1.5)
    first_line_indent = Inches(formatting_profile.get("paragraph_indent_inches", 0.0))
    normal_style = document.styles["Normal"]
    normal_style.font.name = font_name
    normal_style.font.size = Pt(font_size)
    normal_style.paragraph_format.line_spacing = line_spacing
    normal_style.paragraph_format.first_line_indent = first_line_indent

    for style_name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        style = document.styles[style_name]
        style.font.name = font_name
        style.font.size = Pt(size)
        style.paragraph_format.line_spacing = line_spacing
        style.paragraph_format.first_line_indent = Inches(0)

    for style_name in ["List Bullet", "List Number"]:
        style = document.styles[style_name]
        style.font.name = font_name
        style.font.size = Pt(font_size)
        style.paragraph_format.line_spacing = line_spacing

    caption_style = document.styles["Caption"]
    caption_style.font.name = font_name
    caption_style.font.size = Pt(max(font_size - 1, 9))
    caption_style.font.italic = formatting_profile.get("caption_style", "italic") == "italic"
    caption_style.paragraph_format.line_spacing = line_spacing
    caption_style.paragraph_format.first_line_indent = Inches(0)


def _apply_docx_body_paragraph_format(paragraph: Any, formatting_profile: dict[str, Any]) -> None:
    paragraph.paragraph_format.line_spacing = formatting_profile.get("line_spacing", 1.5)
    paragraph.paragraph_format.first_line_indent = Inches(
        formatting_profile.get("paragraph_indent_inches", 0.0)
    )


def _apply_docx_reference_layout(paragraph: Any, formatting_profile: dict[str, Any]) -> None:
    if formatting_profile.get("reference_layout") == "hanging_indent":
        paragraph.paragraph_format.left_indent = Inches(0.3)
        paragraph.paragraph_format.first_line_indent = Inches(-0.3)
    paragraph.paragraph_format.line_spacing = formatting_profile.get("line_spacing", 1.5)


def _heading_text(title: str, numbering: list[int], formatting_profile: dict[str, Any]) -> str:
    if formatting_profile.get("heading_numbering"):
        return f"{'.'.join(str(value) for value in numbering)}. {title}"
    return title


def _table_matrix(table_model: dict[str, Any]) -> list[list[str]]:
    matrix: list[list[str]] = []
    if table_model["headers"]:
        matrix.append(list(table_model["headers"]))
    matrix.extend([list(row) for row in table_model["rows"]])
    if not matrix:
        return []

    column_count = max(len(row) for row in matrix)
    return [
        row + [""] * (column_count - len(row))
        for row in matrix
    ]


def _add_docx_blocks(
    document: Document,
    section: dict[str, Any],
    evidence_state: dict[str, Any],
    formatting_profile: dict[str, Any],
    numbering: list[int],
    heading_level: int,
) -> None:
    for block in section["blocks"]:
        if block["kind"] == "paragraph":
            paragraph = document.add_paragraph(block["text"])
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _apply_docx_body_paragraph_format(paragraph, formatting_profile)
        elif block["kind"] == "bullet_list":
            for item in block["items"]:
                paragraph = document.add_paragraph(item, style="List Bullet")
                paragraph.paragraph_format.line_spacing = formatting_profile.get("line_spacing", 1.5)
        elif block["kind"] == "ordered_list":
            for item in block["items"]:
                paragraph = document.add_paragraph(item, style="List Number")
                paragraph.paragraph_format.line_spacing = formatting_profile.get("line_spacing", 1.5)
        elif block["kind"] == "figure":
            _insert_docx_figure(document, block["figure_id"], evidence_state)
        elif block["kind"] == "table":
            _insert_docx_table(document, block["table_id"], evidence_state, formatting_profile)

    for index, subsection in enumerate(section["subsections"], start=1):
        subsection_numbering = numbering + [index]
        document.add_heading(
            _heading_text(subsection["title"], subsection_numbering, formatting_profile),
            level=min(heading_level, 9),
        )
        _add_docx_blocks(
            document,
            subsection,
            evidence_state,
            formatting_profile,
            subsection_numbering,
            heading_level + 1,
        )


def _insert_docx_figure(document: Document, figure_id: str, evidence_state: dict[str, Any]) -> None:
    figure = lookup_figure(evidence_state, figure_id)
    if not figure:
        document.add_paragraph(f"[Missing figure: {figure_id}]")
        return
    figure_path = Path(figure.get("path", ""))
    if figure_path.exists():
        document.add_picture(str(figure_path), width=Inches(5.8))
    caption = figure_caption_text(figure) or figure_id
    document.add_paragraph(caption, style="Caption")


def _insert_docx_table(
    document: Document,
    table_id: str,
    evidence_state: dict[str, Any],
    formatting_profile: dict[str, Any],
) -> None:
    table = lookup_table(evidence_state, table_id)
    if not table:
        document.add_paragraph(f"[Missing table: {table_id}]")
        return

    table_model = table_display_model(table)
    if table_model["caption"]:
        document.add_paragraph(table_model["caption"], style="Caption")

    matrix = _table_matrix(table_model)
    if not matrix:
        document.add_paragraph("[Table content pending]")
        return

    docx_table = document.add_table(rows=len(matrix), cols=len(matrix[0]))
    if formatting_profile.get("table_style") in {"grid", "compact_grid"}:
        docx_table.style = "Table Grid"

    for row_index, row in enumerate(matrix):
        for column_index, cell_text in enumerate(row):
            paragraph = docx_table.cell(row_index, column_index).paragraphs[0]
            run = paragraph.add_run(cell_text)
            if row_index == 0 and table_model["headers"]:
                run.bold = True
            paragraph.paragraph_format.line_spacing = formatting_profile.get("line_spacing", 1.0)


def export_docx(
    output_path: Path,
    document_model: dict[str, Any],
    formatting_profile: dict[str, Any],
    bibliography: list[str],
    evidence_state: dict[str, Any],
) -> None:
    document = Document()
    _apply_docx_page_layout(document, formatting_profile)
    _configure_docx_styles(document, formatting_profile)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(document_model["title"])
    title_run.bold = True
    title_run.font.size = Pt(16)

    if formatting_profile.get("front_matter") == "title_page_then_abstract":
        document.add_page_break()

    if document_model["abstract_blocks"]:
        document.add_heading("Abstract", level=1)
        for block in document_model["abstract_blocks"]:
            if block["kind"] == "paragraph":
                paragraph = document.add_paragraph(block["text"])
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                _apply_docx_body_paragraph_format(paragraph, formatting_profile)

    if document_model["keywords"]:
        paragraph = document.add_paragraph("Keywords: " + ", ".join(document_model["keywords"]))
        _apply_docx_body_paragraph_format(paragraph, formatting_profile)

    for index, section in enumerate(document_model["sections"], start=1):
        numbering = [index]
        document.add_heading(
            _heading_text(section["title"], numbering, formatting_profile),
            level=1,
        )
        _add_docx_blocks(
            document,
            section,
            evidence_state,
            formatting_profile,
            numbering,
            2,
        )

    if bibliography:
        document.add_page_break()
        document.add_heading("References", level=1)
        for entry in bibliography:
            paragraph = document.add_paragraph(entry)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _apply_docx_reference_layout(paragraph, formatting_profile)

    if document_model["appendices"]:
        document.add_section(WD_SECTION.NEW_PAGE)
        document.add_heading("Appendices", level=1)
        for index, appendix in enumerate(document_model["appendices"], start=1):
            numbering = [index]
            document.add_heading(
                _heading_text(appendix["title"], numbering, formatting_profile),
                level=2,
            )
            _add_docx_blocks(
                document,
                appendix,
                evidence_state,
                formatting_profile,
                numbering,
                3,
            )

    document.save(output_path)


def _pdf_styles(formatting_profile: dict[str, Any]) -> dict[str, ParagraphStyle]:
    font_name = _reportlab_font_name(formatting_profile.get("font_name", "Times New Roman"))
    font_size = formatting_profile.get("font_size_pt", 12)
    spacing = formatting_profile.get("line_spacing", 1.5)
    first_line_indent = formatting_profile.get("paragraph_indent_inches", 0.0) * inch
    styles = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            "PaperTitle",
            parent=styles["Title"],
            fontName=font_name,
            fontSize=16,
            leading=20,
            alignment=1,
            spaceAfter=18,
        ),
        "heading1": ParagraphStyle(
            "PaperHeading1",
            parent=styles["Heading1"],
            fontName=font_name,
            fontSize=14,
            leading=18,
            spaceBefore=14,
            spaceAfter=8,
        ),
        "heading2": ParagraphStyle(
            "PaperHeading2",
            parent=styles["Heading2"],
            fontName=font_name,
            fontSize=12,
            leading=16,
            spaceBefore=10,
            spaceAfter=6,
        ),
        "body": ParagraphStyle(
            "PaperBody",
            parent=styles["BodyText"],
            fontName=font_name,
            fontSize=font_size,
            leading=font_size * spacing,
            spaceAfter=8,
            firstLineIndent=first_line_indent,
        ),
        "caption": ParagraphStyle(
            "PaperCaption",
            parent=styles["Italic"]
            if formatting_profile.get("caption_style", "italic") == "italic"
            else styles["BodyText"],
            fontName=font_name,
            fontSize=10,
            leading=12,
            textColor=colors.black,
            spaceAfter=8,
        ),
        "reference": ParagraphStyle(
            "PaperReference",
            parent=styles["BodyText"],
            fontName=font_name,
            fontSize=font_size,
            leading=font_size * spacing,
            spaceAfter=8,
            leftIndent=0.3 * inch if formatting_profile.get("reference_layout") == "hanging_indent" else 0,
            firstLineIndent=-0.3 * inch if formatting_profile.get("reference_layout") == "hanging_indent" else 0,
        ),
    }


def _insert_pdf_figure(story: list[Any], figure_id: str, evidence_state: dict[str, Any], styles: dict[str, ParagraphStyle]) -> None:
    figure = lookup_figure(evidence_state, figure_id)
    if not figure:
        story.append(Paragraph(f"[Missing figure: {figure_id}]", styles["body"]))
        return
    figure_path = Path(figure.get("path", ""))
    if figure_path.exists():
        story.append(Image(str(figure_path), width=5.5 * inch, height=3.5 * inch))
    story.append(Paragraph(figure_caption_text(figure) or figure_id, styles["caption"]))


def _pdf_table_style(formatting_profile: dict[str, Any]) -> TableStyle:
    compact = formatting_profile.get("table_style") == "compact_grid"
    return TableStyle(
        [
            ("GRID", (0, 0), (-1, -1), 0.5 if compact else 0.8, colors.black),
            ("BACKGROUND", (0, 0), (-1, 0), colors.whitesmoke),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("PADDING", (0, 0), (-1, -1), 4 if compact else 6),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]
    )


def _insert_pdf_table(
    story: list[Any],
    table_id: str,
    evidence_state: dict[str, Any],
    formatting_profile: dict[str, Any],
    styles: dict[str, ParagraphStyle],
) -> None:
    table = lookup_table(evidence_state, table_id)
    if not table:
        story.append(Paragraph(f"[Missing table: {table_id}]", styles["body"]))
        return

    table_model = table_display_model(table)
    if table_model["caption"]:
        story.append(Paragraph(table_model["caption"], styles["caption"]))

    matrix = _table_matrix(table_model)
    if not matrix:
        story.append(Paragraph("[Table content pending]", styles["body"]))
        return

    pdf_table = PdfTable(matrix, hAlign="LEFT")
    pdf_table.setStyle(_pdf_table_style(formatting_profile))
    story.append(pdf_table)
    story.append(Spacer(1, 0.15 * inch))


def _append_pdf_blocks(
    story: list[Any],
    section: dict[str, Any],
    evidence_state: dict[str, Any],
    formatting_profile: dict[str, Any],
    styles: dict[str, ParagraphStyle],
    numbering: list[int],
    heading_style: str,
) -> None:
    for block in section["blocks"]:
        if block["kind"] == "paragraph":
            story.append(Paragraph(block["text"], styles["body"]))
        elif block["kind"] == "bullet_list":
            story.append(
                ListFlowable(
                    [ListItem(Paragraph(item, styles["body"])) for item in block["items"]],
                    bulletType="bullet",
                )
            )
            story.append(Spacer(1, 0.1 * inch))
        elif block["kind"] == "ordered_list":
            story.append(
                ListFlowable(
                    [ListItem(Paragraph(item, styles["body"])) for item in block["items"]],
                    bulletType="1",
                )
            )
            story.append(Spacer(1, 0.1 * inch))
        elif block["kind"] == "figure":
            _insert_pdf_figure(story, block["figure_id"], evidence_state, styles)
        elif block["kind"] == "table":
            _insert_pdf_table(story, block["table_id"], evidence_state, formatting_profile, styles)
    for index, subsection in enumerate(section["subsections"], start=1):
        subsection_numbering = numbering + [index]
        story.append(
            Paragraph(
                _heading_text(subsection["title"], subsection_numbering, formatting_profile),
                styles[heading_style],
            )
        )
        _append_pdf_blocks(
            story,
            subsection,
            evidence_state,
            formatting_profile,
            styles,
            subsection_numbering,
            "heading2",
        )


def export_pdf(
    output_path: Path,
    document_model: dict[str, Any],
    formatting_profile: dict[str, Any],
    bibliography: list[str],
    evidence_state: dict[str, Any],
) -> None:
    page_size = _page_size(formatting_profile)
    margins = formatting_profile.get("margins_inches", {})
    styles = _pdf_styles(formatting_profile)
    document = SimpleDocTemplate(
        str(output_path),
        pagesize=page_size,
        rightMargin=margins.get("right", 1.0) * inch,
        leftMargin=margins.get("left", 1.0) * inch,
        topMargin=margins.get("top", 1.0) * inch,
        bottomMargin=margins.get("bottom", 1.0) * inch,
    )
    story: list[Any] = [Paragraph(document_model["title"], styles["title"])]

    if formatting_profile.get("front_matter") == "title_page_then_abstract":
        story.append(PageBreak())

    if document_model["abstract_blocks"]:
        story.append(Paragraph("Abstract", styles["heading1"]))
        for block in document_model["abstract_blocks"]:
            if block["kind"] == "paragraph":
                story.append(Paragraph(block["text"], styles["body"]))

    if document_model["keywords"]:
        story.append(Paragraph("Keywords: " + ", ".join(document_model["keywords"]), styles["body"]))

    for index, section in enumerate(document_model["sections"], start=1):
        numbering = [index]
        story.append(
            Paragraph(
                _heading_text(section["title"], numbering, formatting_profile),
                styles["heading1"],
            )
        )
        _append_pdf_blocks(
            story,
            section,
            evidence_state,
            formatting_profile,
            styles,
            numbering,
            "heading2",
        )

    if bibliography:
        story.append(Paragraph("References", styles["heading1"]))
        for entry in bibliography:
            story.append(Paragraph(entry, styles["reference"]))

    if document_model["appendices"]:
        story.append(Paragraph("Appendices", styles["heading1"]))
        for index, appendix in enumerate(document_model["appendices"], start=1):
            numbering = [index]
            story.append(
                Paragraph(
                    _heading_text(appendix["title"], numbering, formatting_profile),
                    styles["heading2"],
                )
            )
            _append_pdf_blocks(
                story,
                appendix,
                evidence_state,
                formatting_profile,
                styles,
                numbering,
                "heading2",
            )

    document.build(story)


def export_raw_text(
    output_path: Path,
    document_model: dict[str, Any],
    bibliography: list[str],
    evidence_state: dict[str, Any],
) -> None:
    lines = [document_model["title"], "", "Abstract", ""]
    for block in document_model["abstract_blocks"]:
        if block["kind"] == "paragraph":
            lines.append(block["text"])
        elif block["kind"] in {"bullet_list", "ordered_list"}:
            lines.extend(f"- {item}" for item in block["items"])
    lines.extend(["", "Body", ""])

    def append_section(section: dict[str, Any], depth: int = 0) -> None:
        prefix = "  " * depth
        lines.append(f"{prefix}{section['title']}")
        lines.append("")
        for block in section["blocks"]:
            if block["kind"] == "paragraph":
                lines.append(block["text"])
            elif block["kind"] in {"bullet_list", "ordered_list"}:
                lines.extend(f"{prefix}- {item}" for item in block["items"])
            elif block["kind"] == "figure":
                figure = lookup_figure(evidence_state, block["figure_id"])
                caption = figure_caption_text(figure) if figure else block["figure_id"]
                lines.append(f"{prefix}[Figure] {caption}")
            elif block["kind"] == "table":
                table = lookup_table(evidence_state, block["table_id"])
                if not table:
                    lines.append(f"{prefix}[Missing table: {block['table_id']}]")
                else:
                    table_model = table_display_model(table)
                    if table_model["caption"]:
                        lines.append(f"{prefix}[Table] {table_model['caption']}")
                    if table_model["headers"]:
                        lines.append(f"{prefix}" + " | ".join(table_model["headers"]))
                    for row in table_model["rows"]:
                        lines.append(f"{prefix}" + " | ".join(row))
        lines.append("")
        for subsection in section["subsections"]:
            append_section(subsection, depth + 1)

    for section in document_model["sections"]:
        append_section(section)

    if bibliography:
        lines.extend(["References", ""])
        lines.extend(bibliography)
        lines.append("")

    output_path.write_text("\n".join(lines).strip() + "\n", encoding="utf-8")


def export_project(root: Path) -> dict[str, str]:
    states = load_all_states(root)
    project_state = states["project"]
    workflow_state = states["workflow"]
    manuscript_state = states["manuscript"]
    sources_state = states["sources"]
    evidence_state = states["evidence"]
    exports_state = states["exports"]

    validation_report = validate_project_state(
        project_state, workflow_state, manuscript_state, sources_state, evidence_state
    )
    exports_state["last_validation"] = validation_report
    if not validation_report["valid"]:
        save_state(root, "exports.json", exports_state)
        append_trace(
            root,
            {
                "event": "export_blocked",
                "reason": "validation_failed",
                "finding_count": len(validation_report["findings"]),
            },
        )
        raise ValueError("Export blocked by validation failures. Run `spw validate` for details.")

    citation_order = build_citation_order(manuscript_state, evidence_state)
    sources_by_id = source_index(sources_state)
    citation_style = project_state["profiles"]["citation"]
    document_model = render_document_model(
        project_state, manuscript_state, sources_by_id, citation_style, citation_order
    )
    bibliography = bibliography_entries(sources_state, citation_style, citation_order)

    outputs: dict[str, str] = {}
    deliverables_dir = root / "deliverables"
    formatting_profile = resolve_effective_formatting_profile(project_state)
    if project_state.get("profile_resolution", {}).get("effective_formatting") != formatting_profile:
        project_state.setdefault("profile_resolution", {})["effective_formatting"] = formatting_profile
        save_state(root, "project.json", project_state)

    if project_state["output_mode"] in {"raw", "both"}:
        raw_path = deliverables_dir / "final_paper.txt"
        export_raw_text(raw_path, document_model, bibliography, evidence_state)
        outputs["raw_text"] = str(raw_path)
        exports_state["raw_text"]["status"] = "completed"

    if project_state["output_mode"] in {"formatted", "both"}:
        docx_path = deliverables_dir / "final_paper.docx"
        pdf_path = deliverables_dir / "final_paper.pdf"
        export_docx(docx_path, document_model, formatting_profile, bibliography, evidence_state)
        export_pdf(pdf_path, document_model, formatting_profile, bibliography, evidence_state)
        outputs["docx"] = str(docx_path)
        outputs["pdf"] = str(pdf_path)
        exports_state["docx"]["status"] = "completed"
        exports_state["pdf"]["status"] = "completed"

    save_state(root, "exports.json", exports_state)
    append_trace(root, {"event": "export_completed", "outputs": outputs})
    return outputs
