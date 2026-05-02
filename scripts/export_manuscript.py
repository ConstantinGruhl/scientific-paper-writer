from __future__ import annotations

import argparse
import json
import re
import zipfile
from pathlib import Path
from xml.sax.saxutils import escape


def normalize_text(text: str) -> str:
    return text.replace("\r\n", "\n").replace("\r", "\n")


def read_manifest(workspace_dir: Path) -> dict:
    manifest_path = workspace_dir / "manifest.json"
    if not manifest_path.exists():
        return {"output_format": "docx", "output_file": "paper.docx"}
    return json.loads(manifest_path.read_text(encoding="utf-8"))


def choose_source(workspace_dir: Path) -> Path:
    final_path = workspace_dir / "final_paper.md"
    draft_path = workspace_dir / "draft.md"
    if final_path.exists() and final_path.read_text(encoding="utf-8").strip():
        return final_path
    return draft_path


def markdown_blocks(text: str) -> list[tuple[str, str]]:
    blocks: list[tuple[str, str]] = []
    paragraph: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph
        if paragraph:
            blocks.append(("paragraph", " ".join(line.strip() for line in paragraph).strip()))
            paragraph = []

    for raw_line in normalize_text(text).split("\n"):
        line = raw_line.rstrip()
        stripped = line.strip()

        if not stripped:
            flush_paragraph()
            continue

        if stripped.startswith("#"):
            flush_paragraph()
            level = len(stripped) - len(stripped.lstrip("#"))
            content = stripped[level:].strip()
            blocks.append((f"heading{min(level, 3)}", content))
            continue

        bullet_match = re.match(r"^[-*]\s+(.*)$", stripped)
        if bullet_match:
            flush_paragraph()
            blocks.append(("bullet", bullet_match.group(1).strip()))
            continue

        paragraph.append(stripped)

    flush_paragraph()
    return blocks


def markdown_to_latex(text: str) -> str:
    title = "Paper"
    lines = [
        r"\documentclass[12pt]{article}",
        r"\usepackage[utf8]{inputenc}",
        r"\usepackage[T1]{fontenc}",
        r"\usepackage{enumitem}",
        r"\usepackage{geometry}",
        r"\geometry{margin=1in}",
        r"\begin{document}",
    ]
    in_list = False

    for kind, content in markdown_blocks(text):
        safe = (
            content.replace("\\", r"\textbackslash{}")
            .replace("&", r"\&")
            .replace("%", r"\%")
            .replace("$", r"\$")
            .replace("#", r"\#")
            .replace("_", r"\_")
            .replace("{", r"\{")
            .replace("}", r"\}")
        )

        if kind.startswith("heading"):
            if in_list:
                lines.append(r"\end{itemize}")
                in_list = False
            if kind == "heading1":
                title = content or title
                lines.append(rf"\section*{{{safe}}}")
            elif kind == "heading2":
                lines.append(rf"\subsection*{{{safe}}}")
            else:
                lines.append(rf"\subsubsection*{{{safe}}}")
        elif kind == "bullet":
            if not in_list:
                lines.append(r"\begin{itemize}[leftmargin=1.5em]")
                in_list = True
            lines.append(rf"\item {safe}")
        else:
            if in_list:
                lines.append(r"\end{itemize}")
                in_list = False
            lines.append(safe + "\n")

    if in_list:
        lines.append(r"\end{itemize}")

    lines.append(r"\end{document}")
    return "\n".join(lines) + "\n"


def build_docx_paragraph(text: str, style: str | None = None, bullet: bool = False) -> str:
    ppr = ""
    if style:
        ppr = f"<w:pPr><w:pStyle w:val=\"{style}\"/></w:pPr>"
    prefix = "• " if bullet else ""
    return (
        "<w:p>"
        f"{ppr}"
        "<w:r><w:t xml:space=\"preserve\">"
        f"{escape(prefix + text)}"
        "</w:t></w:r>"
        "</w:p>"
    )


def markdown_to_docx_bytes(text: str) -> bytes:
    try:
        from docx import Document  # type: ignore

        document = Document()
        for kind, content in markdown_blocks(text):
            if kind == "heading1":
                document.add_heading(content, level=1)
            elif kind == "heading2":
                document.add_heading(content, level=2)
            elif kind == "heading3":
                document.add_heading(content, level=3)
            elif kind == "bullet":
                document.add_paragraph(content, style="List Bullet")
            else:
                document.add_paragraph(content)

        from io import BytesIO

        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue()
    except Exception:
        body_parts: list[str] = []
        for kind, content in markdown_blocks(text):
            if kind == "heading1":
                body_parts.append(build_docx_paragraph(content, style="Heading1"))
            elif kind == "heading2":
                body_parts.append(build_docx_paragraph(content, style="Heading2"))
            elif kind == "heading3":
                body_parts.append(build_docx_paragraph(content, style="Heading3"))
            elif kind == "bullet":
                body_parts.append(build_docx_paragraph(content, bullet=True))
            else:
                body_parts.append(build_docx_paragraph(content))

        document_xml = (
            "<?xml version=\"1.0\" encoding=\"UTF-8\" standalone=\"yes\"?>"
            "<w:document xmlns:wpc=\"http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas\" "
            "xmlns:mc=\"http://schemas.openxmlformats.org/markup-compatibility/2006\" "
            "xmlns:o=\"urn:schemas-microsoft-com:office:office\" "
            "xmlns:r=\"http://schemas.openxmlformats.org/officeDocument/2006/relationships\" "
            "xmlns:m=\"http://schemas.openxmlformats.org/officeDocument/2006/math\" "
            "xmlns:v=\"urn:schemas-microsoft-com:vml\" "
            "xmlns:wp14=\"http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing\" "
            "xmlns:wp=\"http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing\" "
            "xmlns:w10=\"urn:schemas-microsoft-com:office:word\" "
            "xmlns:w=\"http://schemas.openxmlformats.org/wordprocessingml/2006/main\" "
            "xmlns:w14=\"http://schemas.microsoft.com/office/word/2010/wordml\" "
            "xmlns:w15=\"http://schemas.microsoft.com/office/word/2012/wordml\" "
            "xmlns:wpg=\"http://schemas.microsoft.com/office/word/2010/wordprocessingGroup\" "
            "xmlns:wpi=\"http://schemas.microsoft.com/office/word/2010/wordprocessingInk\" "
            "xmlns:wne=\"http://schemas.microsoft.com/office/word/2006/wordml\" "
            "xmlns:wps=\"http://schemas.microsoft.com/office/word/2010/wordprocessingShape\" "
            "mc:Ignorable=\"w14 w15 wp14\">"
            "<w:body>"
            + "".join(body_parts)
            + "<w:sectPr><w:pgSz w:w=\"12240\" w:h=\"15840\"/><w:pgMar w:top=\"1440\" w:right=\"1440\" w:bottom=\"1440\" w:left=\"1440\" w:header=\"708\" w:footer=\"708\" w:gutter=\"0\"/></w:sectPr>"
            "</w:body></w:document>"
        )

        content_types = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
  <Override PartName="/word/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/>
</Types>
"""

        rels = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>
"""

        document_rels = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"/>
"""

        styles = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:style w:type="paragraph" w:default="1" w:styleId="Normal">
    <w:name w:val="Normal"/>
  </w:style>
  <w:style w:type="paragraph" w:styleId="Heading1">
    <w:name w:val="heading 1"/>
    <w:basedOn w:val="Normal"/>
    <w:qFormat/>
  </w:style>
  <w:style w:type="paragraph" w:styleId="Heading2">
    <w:name w:val="heading 2"/>
    <w:basedOn w:val="Normal"/>
    <w:qFormat/>
  </w:style>
  <w:style w:type="paragraph" w:styleId="Heading3">
    <w:name w:val="heading 3"/>
    <w:basedOn w:val="Normal"/>
    <w:qFormat/>
  </w:style>
  <w:style w:type="paragraph" w:styleId="ListParagraph">
    <w:name w:val="List Paragraph"/>
    <w:basedOn w:val="Normal"/>
  </w:style>
</w:styles>
"""

        from io import BytesIO

        buffer = BytesIO()
        with zipfile.ZipFile(buffer, "w", compression=zipfile.ZIP_DEFLATED) as archive:
            archive.writestr("[Content_Types].xml", content_types)
            archive.writestr("_rels/.rels", rels)
            archive.writestr("word/document.xml", document_xml)
            archive.writestr("word/_rels/document.xml.rels", document_rels)
            archive.writestr("word/styles.xml", styles)
        return buffer.getvalue()


def export_workspace(folder: Path, output_format: str | None = None) -> Path:
    workspace_dir = folder / ".paper_writer"
    if not workspace_dir.exists():
        raise SystemExit(f"Workspace not found: {workspace_dir}")

    manifest = read_manifest(workspace_dir)
    chosen_format = output_format or manifest.get("output_format", "docx")
    source_path = choose_source(workspace_dir)
    source_text = source_path.read_text(encoding="utf-8")

    if chosen_format == "latex":
        output_path = folder / "paper.tex"
        output_path.write_text(markdown_to_latex(source_text), encoding="utf-8")
        return output_path

    output_path = folder / "paper.docx"
    output_path.write_bytes(markdown_to_docx_bytes(source_text))
    return output_path


def main() -> int:
    parser = argparse.ArgumentParser(
        description="Export the final paper from a folder-native workspace."
    )
    parser.add_argument(
        "target",
        nargs="?",
        default=".",
        help="Folder containing .paper_writer/",
    )
    parser.add_argument(
        "--format",
        choices=["latex", "docx"],
        help="Override the manifest output format",
    )
    args = parser.parse_args()

    target_dir = Path(args.target).resolve()
    output_path = export_workspace(target_dir, output_format=args.format)
    print(output_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
